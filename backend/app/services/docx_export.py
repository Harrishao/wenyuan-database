import re
from io import BytesIO
from uuid import UUID

from docx import Document as DocxDocument

from app.domain.models import Document, Report, ReportSection


def render_report_docx(
    report: Report,
    sections: list[ReportSection],
    references: list[Document] | None = None,
    citation_numbers: dict[UUID, dict[str, int]] | None = None,
) -> bytes:
    document = DocxDocument()
    document.add_heading(report.title, level=0)
    for section in sorted(sections, key=lambda item: item.position):
        document.add_heading(section.title, level=1)
        content = section.content_markdown
        if citation_numbers:
            marker_map = citation_numbers.get(section.id, {})
            content = re.sub(
                r"\[\d+\]",
                lambda match, markers=marker_map: (
                    f"[{markers[match.group(0)]}]" if match.group(0) in markers else match.group(0)
                ),
                content,
            )
        for block in content.split("\n\n"):
            text = block.strip()
            if not text:
                continue
            if text.startswith("### "):
                document.add_heading(text[4:], level=2)
            else:
                document.add_paragraph(text)
    if references:
        document.add_heading("参考文献", level=1)
        for index, reference in enumerate(references, start=1):
            title = reference.publication_title or reference.original_filename
            parts = [f"[{index}] {reference.author}. {title}"]
            if reference.source:
                parts.append(reference.source)
            if reference.publication_year:
                parts.append(str(reference.publication_year))
            document.add_paragraph(". ".join(parts) + ".")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
