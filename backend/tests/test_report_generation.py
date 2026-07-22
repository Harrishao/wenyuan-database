import json
from uuid import uuid4

import pytest
from docx import Document as DocxDocument

from app.adapters.report_llm import LocalEvidenceDraftLlm
from app.domain.enums import ReportStatus
from app.domain.models import Report, ReportSection
from app.ports.llm import ChatMessage, ChatOptions
from app.services.docx_export import render_report_docx
from app.services.report_generation import validate_citation_markers


@pytest.mark.asyncio
async def test_local_report_draft_uses_only_supplied_evidence() -> None:
    llm = LocalEvidenceDraftLlm()
    prompt = json.dumps(
        {
            "topic": "光伏功率预测",
            "section_title": "研究背景",
            "evidence": [
                {"content": "短期预测可为电网调度提供决策依据。"},
                {"content": "时序模型能够描述出力变化规律。"},
            ],
        },
        ensure_ascii=False,
    )
    result = await llm.chat(
        [ChatMessage(role="user", content=prompt)],
        ChatOptions(),
    )
    assert "[1]" in result
    assert "[2]" in result
    assert "[3]" not in result


def test_invalid_citation_markers_are_removed() -> None:
    content, markers = validate_citation_markers("结论[1]，无效来源[7]。", 2)
    assert content == "结论[1]，无效来源。"
    assert markers == [1]


def test_docx_export_can_be_opened(tmp_path) -> None:
    report = Report(
        id=uuid4(),
        owner_id=uuid4(),
        knowledge_base_id=uuid4(),
        template_version_id=uuid4(),
        title="测试报告",
        status=ReportStatus.READY,
        current_version=1,
    )
    section = ReportSection(
        report_version_id=uuid4(),
        key="background",
        title="研究背景",
        position=1,
        content_markdown="证据支撑的正文。[1]",
    )
    output = tmp_path / "report.docx"
    output.write_bytes(render_report_docx(report, [section]))
    reopened = DocxDocument(output)
    assert any("测试报告" in paragraph.text for paragraph in reopened.paragraphs)
    assert any("证据支撑的正文" in paragraph.text for paragraph in reopened.paragraphs)
