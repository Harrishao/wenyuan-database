from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from pydantic import ValidationError

from app.domain.models import Document, Report, ReportSection
from app.schemas.admin import AdminTemplateSectionInput
from app.schemas.mvp5 import TemplateVersionInput
from app.services.docx_export import render_report_docx


def test_template_version_requires_contiguous_unique_section_positions() -> None:
    with pytest.raises(ValidationError):
        TemplateVersionInput(
            system_prompt="test",
            settings={"top_k": 4},
            sections=[
                AdminTemplateSectionInput(
                    key="first",
                    title="第一章",
                    position=1,
                    instructions="",
                ),
                AdminTemplateSectionInput(
                    key="second",
                    title="第二章",
                    position=3,
                    instructions="",
                ),
            ],
        )


def test_docx_references_are_stable_and_section_markers_are_remapped() -> None:
    report = Report(
        owner_id=uuid4(),
        knowledge_base_id=uuid4(),
        template_version_id=uuid4(),
        title="测试报告",
    )
    section = ReportSection(
        report_version_id=uuid4(),
        key="intro",
        title="引言",
        position=1,
        content_markdown="证据甲[1]，证据乙[2]。",
    )
    first = Document(
        knowledge_base_id=uuid4(),
        uploaded_by=uuid4(),
        original_filename="甲.txt",
        storage_key="first",
        mime_type="text/plain",
        file_size=1,
        sha256="a" * 64,
        author="甲作者",
        publication_title="甲文献",
        publication_year=2024,
        source="甲期刊",
    )
    second = Document(
        knowledge_base_id=uuid4(),
        uploaded_by=uuid4(),
        original_filename="乙.txt",
        storage_key="second",
        mime_type="text/plain",
        file_size=1,
        sha256="b" * 64,
        author="乙作者",
        publication_title="乙文献",
        publication_year=2025,
        source="乙期刊",
    )
    payload = render_report_docx(
        report,
        [section],
        [second, first],
        {section.id: {"[1]": 2, "[2]": 1}},
    )
    document = DocxDocument(BytesIO(payload))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "证据甲[2]，证据乙[1]。" in paragraphs
    assert "[1] 乙作者. 乙文献. 乙期刊. 2025." in paragraphs
    assert "[2] 甲作者. 甲文献. 甲期刊. 2024." in paragraphs
